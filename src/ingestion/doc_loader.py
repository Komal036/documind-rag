"""
DocuMind Document Loader
-------------------------
Handles loading raw documents (PDF, TXT, DOCX, MD) and converting them
into a normalised list of LangChain Document objects with rich metadata.

Supported formats:
  • .pdf  — via PyMuPDF (fast) with pdfplumber fallback
  • .txt  — plain text
  • .md   — Markdown (treated as plain text)
  • .docx — python-docx
"""
from __future__ import annotations

import datetime
from pathlib import Path
from typing import List

from langchain_core.documents import Document

from src.utils.exceptions import (
    DocumentLoadError,
    DocumentParsingError,
    UnsupportedFileTypeError,
)
from src.utils.file_utils import compute_file_hash, get_file_metadata
from src.utils.logger import get_logger

logger = get_logger(__name__)


# ── Public API ────────────────────────────────────────────────────────

def load_document(file_path: Path) -> List[Document]:
    """
    Load a single document and return a list of LangChain Documents.

    Each Document has metadata:
        source, filename, extension, size_bytes, sha256,
        page (for PDFs), loaded_at

    Args:
        file_path: Absolute or relative path to the document.

    Returns:
        List[Document] — one entry per page (PDF) or one entry (others).

    Raises:
        UnsupportedFileTypeError: If the file type is not supported.
        DocumentLoadError: If the file cannot be opened.
        DocumentParsingError: If text extraction fails.
    """
    file_path = Path(file_path).resolve()
    ext = file_path.suffix.lower()

    logger.info("Loading document", path=str(file_path), ext=ext)

    loaders = {
        ".pdf": _load_pdf,
        ".txt": _load_text,
        ".md": _load_text,
        ".docx": _load_docx,
    }

    if ext not in loaders:
        raise UnsupportedFileTypeError(
            f"Unsupported file type: {ext}",
            details={"supported": list(loaders.keys()), "received": ext},
        )

    try:
        docs = loaders[ext](file_path)
    except (UnsupportedFileTypeError, DocumentLoadError, DocumentParsingError):
        raise
    except Exception as exc:
        raise DocumentLoadError(
            f"Failed to load '{file_path.name}': {exc}",
            details={"path": str(file_path)},
        ) from exc

    # Enrich metadata on every page/doc
    meta = get_file_metadata(file_path)
    loaded_at = datetime.datetime.utcnow().isoformat()
    for doc in docs:
        doc.metadata.update(
            {
                "source": str(file_path),
                "filename": meta["filename"],
                "extension": meta["extension"],
                "size_bytes": meta["size_bytes"],
                "sha256": meta["sha256"],
                "loaded_at": loaded_at,
            }
        )

    logger.info(
        "Document loaded",
        filename=file_path.name,
        pages=len(docs),
        sha256=meta["sha256"][:8],
    )
    return docs


# ── Private loaders ───────────────────────────────────────────────────

def _load_pdf(file_path: Path) -> List[Document]:
    """Extract text from PDF using PyMuPDF; falls back to pdfplumber."""
    try:
        return _load_pdf_pymupdf(file_path)
    except Exception as primary_exc:
        logger.warning(
            "PyMuPDF failed, falling back to pdfplumber",
            error=str(primary_exc),
        )
        try:
            return _load_pdf_pdfplumber(file_path)
        except Exception as fallback_exc:
            raise DocumentParsingError(
                f"Both PDF parsers failed for '{file_path.name}'",
                details={"pymupdf": str(primary_exc), "pdfplumber": str(fallback_exc)},
            ) from fallback_exc


def _load_pdf_pymupdf(file_path: Path) -> List[Document]:
    import fitz  # PyMuPDF

    docs: List[Document] = []
    with fitz.open(str(file_path)) as pdf:
        for page_num, page in enumerate(pdf, start=1):
            text = page.get_text("text").strip()
            if text:
                docs.append(
                    Document(
                        page_content=text,
                        metadata={"page": page_num, "total_pages": len(pdf)},
                    )
                )
    if not docs:
        raise DocumentParsingError(f"No extractable text in '{file_path.name}'")
    return docs


def _load_pdf_pdfplumber(file_path: Path) -> List[Document]:
    import pdfplumber

    docs: List[Document] = []
    with pdfplumber.open(str(file_path)) as pdf:
        total = len(pdf.pages)
        for page_num, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                docs.append(
                    Document(
                        page_content=text,
                        metadata={"page": page_num, "total_pages": total},
                    )
                )
    if not docs:
        raise DocumentParsingError(f"No extractable text in '{file_path.name}'")
    return docs


def _load_text(file_path: Path) -> List[Document]:
    """Load plain text / markdown as a single Document."""
    try:
        text = file_path.read_text(encoding="utf-8").strip()
    except UnicodeDecodeError:
        text = file_path.read_text(encoding="latin-1").strip()

    if not text:
        raise DocumentParsingError(f"Empty file: '{file_path.name}'")

    return [Document(page_content=text, metadata={"page": 1, "total_pages": 1})]


def _load_docx(file_path: Path) -> List[Document]:
    """Extract text from a DOCX file paragraph by paragraph, returned as one Document."""
    try:
        from docx import Document as DocxDocument  # python-docx
    except ImportError as exc:
        raise DocumentLoadError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from exc

    doc = DocxDocument(str(file_path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not paragraphs:
        raise DocumentParsingError(f"No text found in '{file_path.name}'")

    full_text = "\n\n".join(paragraphs)
    return [Document(page_content=full_text, metadata={"page": 1, "total_pages": 1})]
